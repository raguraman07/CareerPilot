import os
import logging
import PyPDF2
import docx

logger = logging.getLogger(__name__)

def parse_pdf(filepath):
    """
    Extract text and page count from a PDF file using PyPDF2.
    Handles encrypted files, per-page extraction exceptions, and empty text fallbacks.
    """
    text = []
    pages = 0
    try:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            # Decrypt if encrypted with blank password
            if getattr(reader, 'is_encrypted', False):
                try:
                    reader.decrypt('')
                except Exception as decrypt_err:
                    logger.warning(f"PDF {filepath} is encrypted: {decrypt_err}")

            pages = len(reader.pages)
            for page_num in range(pages):
                try:
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text.append(page_text.strip())
                except Exception as page_err:
                    logger.warning(f"Could not extract text from page {page_num} of {filepath}: {page_err}")

    except Exception as e:
        logger.error(f"Error parsing PDF {filepath}: {e}")
        raise ValueError(f"Failed to parse PDF file: {str(e)}")
        
    extracted_text = "\n\n".join(text).strip()
    if not extracted_text:
        extracted_text = "[Notice: No selectable text content could be extracted from this PDF document. It may contain scanned image pages.]"

    return {
        "text": extracted_text,
        "pages": pages if pages > 0 else 1,
        "file_type": "pdf"
    }

def parse_docx(filepath):
    """
    Extract text and page count estimate from a DOCX file using python-docx.
    """
    text = []
    try:
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        
        # Also extract tables text
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
                    
    except Exception as e:
        logger.error(f"Error parsing DOCX {filepath}: {e}")
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")
        
    extracted_text = "\n\n".join(text).strip()
    if not extracted_text:
        extracted_text = "[Notice: No readable text content found in DOCX file.]"

    return {
        "text": extracted_text,
        "pages": 1,
        "file_type": "docx"
    }

def parse_resume(filepath):
    """
    Auto-detects format from extension and extracts text.
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
        
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.pdf':
        return parse_pdf(filepath)
    elif ext in ['.docx', '.doc']:
        return parse_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")